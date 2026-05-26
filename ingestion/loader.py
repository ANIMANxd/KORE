"""
Data loader supporting both local exports and live API ingestion.

Local exports are parsed directly from JSON.
Live ingestion uses the Slack SDK (installed with llama-index-readers-slack).
"""

import json
import os
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Literal

from dotenv import load_dotenv
from llama_index.core.schema import Document

load_dotenv()


def _load_json(path: Path) -> Any:
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def _is_valid_message(msg: dict) -> bool:
    if msg.get("type") != "message":
        return False
    if msg.get("subtype") in ("channel_join", "channel_leave", "bot_message"):
        return False
    text = msg.get("text", "").strip()
    if not text:
        return False
    if msg.get("bot_id") or msg.get("user", "").startswith("B"):
        return False
    return True


def _parse_ts(ts: str) -> datetime:
    return datetime.fromtimestamp(float(ts))


def _load_slack_export(
    export_path: str,
    channels: list[str] | None = None,
) -> list[Document]:
    """Load a local Slack export directory into LlamaIndex Documents."""
    root = Path(export_path)
    if not root.exists():
        raise FileNotFoundError(f"Slack export directory not found: {root}")

    users_data = _load_json(root / "users.json") if (root / "users.json").exists() else []
    user_map = {u["id"]: u.get("name", u.get("real_name", u["id"])) for u in users_data}

    channels_data = (
        _load_json(root / "channels.json") if (root / "channels.json").exists() else []
    )
    channel_map = {c["id"]: c["name"] for c in channels_data}

    docs: list[Document] = []

    for ch_dir in root.iterdir():
        if not ch_dir.is_dir():
            continue
        ch_name = ch_dir.name
        if channels and ch_name not in channels:
            continue

        channel_id = None
        for c in channels_data:
            if c["name"] == ch_name:
                channel_id = c["id"]
                break

        daily_files = sorted(ch_dir.glob("*.json"))
        for day_file in daily_files:
            messages = _load_json(day_file)
            if not isinstance(messages, list):
                continue

            for msg in messages:
                if not _is_valid_message(msg):
                    continue

                user_id = msg.get("user", "")
                ts = msg.get("ts", "")
                thread_ts = msg.get("thread_ts")
                text = msg["text"].strip()
                user_name = user_map.get(user_id, user_id)
                dt = _parse_ts(ts)

                metadata = {
                    "source_type": "slack",
                    "channel": f"#{ch_name}",
                    "channel_id": channel_id,
                    "timestamp": dt.isoformat() + "Z",
                    "user_id": user_id,
                    "user_name": user_name,
                    "message_id": ts,
                    "thread_ts": thread_ts,
                    "date_file": day_file.stem,
                }

                doc = Document(
                    text=text,
                    metadata=metadata,
                    id_=f"slack_{ch_name}_{ts.replace('.', '_')}",
                )
                docs.append(doc)

    return docs


def _load_slack_live() -> list[Document]:
    """Load messages from live Slack API using the Slack SDK.

    Requires environment variables:
      SLACK_BOT_TOKEN - Bot User OAuth Token (xoxb-...)
      SLACK_CHANNEL_IDS - Comma-separated list of channel IDs (e.g. C01234567,C09876543)
      SLACK_EARLIEST_DATE - ISO date string, optional (default: 90 days ago)

    The bot must be invited to each channel and have these scopes:
      channels:history, channels:read, users:read
    """
    from slack_sdk import WebClient
    from slack_sdk.errors import SlackApiError

    token = os.getenv("SLACK_BOT_TOKEN")
    if not token:
        raise RuntimeError(
            "SLACK_BOT_TOKEN not set. Add it to your .env file.\n"
            "Create a Slack app at https://api.slack.com/apps, install it to your workspace,\n"
            "and copy the Bot User OAuth Token (starts with xoxb-).\n"
            "Required bot token scopes: channels:history, channels:read, users:read"
        )

    channel_ids_raw = os.getenv("SLACK_CHANNEL_IDS")
    if not channel_ids_raw:
        raise RuntimeError(
            "SLACK_CHANNEL_IDS not set. Add it to your .env file as a comma-separated list,\n"
            "e.g. SLACK_CHANNEL_IDS=C01234567,C09876543\n"
            "Find channel IDs in Slack: right-click channel name → Copy link → ID is the last segment."
        )
    channel_ids = [c.strip() for c in channel_ids_raw.split(",") if c.strip()]

    earliest_date_str = os.getenv("SLACK_EARLIEST_DATE")
    if earliest_date_str:
        earliest_dt = datetime.fromisoformat(earliest_date_str)
        oldest_ts = str(earliest_dt.timestamp())
    else:
        from datetime import timedelta
        oldest_dt = datetime.now(timezone.utc) - timedelta(days=90)
        oldest_ts = str(oldest_dt.timestamp())

    client = WebClient(token=token)

    # Fetch users for name mapping
    try:
        users_result = client.users_list()
        user_map = {
            u["id"]: u.get("name", u.get("real_name", u["id"]))
            for u in users_result.get("members", [])
        }
    except SlackApiError:
        user_map = {}

    docs: list[Document] = []

    for channel_id in channel_ids:
        try:
            # Get channel info for name
            channel_info = client.conversations_info(channel=channel_id)
            channel_name = channel_info.get("channel", {}).get("name", channel_id)
        except SlackApiError:
            channel_name = channel_id

        # Fetch messages
        next_cursor = None
        while True:
            kwargs = {
                "channel": channel_id,
                "cursor": next_cursor,
                "oldest": oldest_ts,
                "limit": 200,
            }
            result = client.conversations_history(**kwargs)
            messages = result.get("messages", [])

            for msg in messages:
                text = msg.get("text", "").strip()
                if not text:
                    continue
                if msg.get("subtype") in ("channel_join", "channel_leave", "bot_message"):
                    continue
                if msg.get("bot_id") or msg.get("user", "").startswith("B"):
                    continue

                user_id = msg.get("user", "")
                ts = msg.get("ts", "")
                thread_ts = msg.get("thread_ts")
                user_name = user_map.get(user_id, user_id)
                dt = datetime.fromtimestamp(float(ts), tz=timezone.utc)

                metadata = {
                    "source_type": "slack",
                    "channel": f"#{channel_name}",
                    "channel_id": channel_id,
                    "timestamp": dt.isoformat(),
                    "user_id": user_id,
                    "user_name": user_name,
                    "message_id": ts,
                    "thread_ts": thread_ts,
                }

                doc = Document(
                    text=text,
                    metadata=metadata,
                    id_=f"slack_{channel_name}_{ts.replace('.', '_')}",
                )
                docs.append(doc)

            if not result.get("has_more"):
                break
            next_cursor = result.get("response_metadata", {}).get("next_cursor")

    return docs


def _load_teams() -> list[Document]:
    """Load messages from Microsoft Teams via Microsoft Graph API.

    Requires environment variables:
      TEAMS_CLIENT_ID     - Azure App Registration Client ID
      TEAMS_CLIENT_SECRET - Azure App Registration Client Secret
      TEAMS_TENANT_ID     - Azure Tenant ID
      TEAMS_TEAM_IDS      - Comma-separated list of Team IDs to ingest

    The Azure app must have these API permissions with admin consent:
      ChannelMessage.Read.All, Chat.Read (for chat messages)
    """
    import requests

    import click as _click
    from rich.console import Console
    from rich.panel import Panel
    from rich.progress import (
        Progress,
        SpinnerColumn,
        TextColumn,
    )

    console = Console()

    client_id = os.getenv("TEAMS_CLIENT_ID")
    client_secret = os.getenv("TEAMS_CLIENT_SECRET")
    tenant_id = os.getenv("TEAMS_TENANT_ID")
    team_ids_raw = os.getenv("TEAMS_TEAM_IDS")

    if not all((client_id, client_secret, tenant_id, team_ids_raw)):
        console.print(Panel("[bold]Microsoft Teams Setup[/bold]", style="bold blue"))
        print("To connect Microsoft Teams:")
        print("1. Go to portal.azure.com → App registrations → New registration")
        print("2. Add API permissions: ChannelMessage.Read.All, Chat.Read")
        print("3. Create a client secret and copy the values below")

        if not client_id:
            client_id = _click.prompt("Enter Azure App Client ID:")
        if not client_secret:
            client_secret = _click.prompt(
                "Enter Azure App Client Secret:", hide_input=True
            )
        if not tenant_id:
            tenant_id = _click.prompt("Enter Azure Tenant ID:")
        if not team_ids_raw:
            team_ids_raw = _click.prompt(
                "Enter Team IDs to ingest (comma-separated):"
            )

        _write_teams_env(client_id, client_secret, tenant_id, team_ids_raw)

    team_ids = [t.strip() for t in team_ids_raw.split(",") if t.strip()]

    # OAuth2 client credentials flow
    token_url = f"https://login.microsoftonline.com/{tenant_id}/oauth2/v2.0/token"
    token_resp = requests.post(
        token_url,
        data={
            "client_id": client_id,
            "client_secret": client_secret,
            "scope": "https://graph.microsoft.com/.default",
            "grant_type": "client_credentials",
        },
        timeout=30,
    )
    token_resp.raise_for_status()
    access_token = token_resp.json()["access_token"]

    headers = {
        "Authorization": f"Bearer {access_token}",
        "Content-Type": "application/json",
    }

    docs: list[Document] = []

    with Progress(
        SpinnerColumn(),
        TextColumn("[cyan]{task.description}"),
        console=console,
    ) as progress:
        for team_id in team_ids:
            # Get team info
            team_info_resp = requests.get(
                f"https://graph.microsoft.com/v1.0/teams/{team_id}",
                headers=headers,
                timeout=30,
            )
            if team_info_resp.status_code == 200:
                team_name = team_info_resp.json().get("displayName", team_id)
            else:
                team_name = team_id

            # Get channels in the team
            channels_resp = requests.get(
                f"https://graph.microsoft.com/v1.0/teams/{team_id}/channels",
                headers=headers,
                timeout=30,
            )
            if channels_resp.status_code != 200:
                console.print(
                    f"[red]Failed to fetch channels for team {team_id}: "
                    f"{channels_resp.status_code}[/red]"
                )
                continue

            channels = channels_resp.json().get("value", [])

            for channel in channels:
                channel_id = channel.get("id")
                channel_name = channel.get("displayName", channel_id)

                task = progress.add_task(
                    f"Fetching {team_name}/{channel_name}...",
                    total=None,
                )

                # Fetch messages in channel
                messages_url = (
                    f"https://graph.microsoft.com/v1.0/teams/{team_id}"
                    f"/channels/{channel_id}/messages"
                )
                while messages_url:
                    msg_resp = requests.get(
                        messages_url,
                        headers=headers,
                        timeout=30,
                    )
                    if msg_resp.status_code != 200:
                        break

                    msg_data = msg_resp.json()
                    messages = msg_data.get("value", [])

                    for msg in messages:
                        body = msg.get("body", {})
                        content = body.get("content", "").strip()
                        if not content:
                            continue

                        # Skip system / event messages where possible
                        msg_type = msg.get("messageType", "")
                        if msg_type and msg_type != "message":
                            continue

                        from_user = msg.get("from", {})
                        user_info = from_user.get("user", {}) if isinstance(from_user, dict) else {}
                        user_id = user_info.get("id", "")
                        user_name = user_info.get("displayName", user_id)

                        created = msg.get("createdDateTime", "")

                        metadata = {
                            "source_type": "teams",
                            "team_id": team_id,
                            "team_name": team_name,
                            "channel_id": channel_id,
                            "channel_name": channel_name,
                            "timestamp": created,
                            "user_id": user_id,
                            "user_name": user_name,
                            "message_id": msg.get("id", ""),
                        }

                        doc = Document(
                            text=content,
                            metadata=metadata,
                            id_=f"teams_{team_id}_{channel_id}_{msg.get('id', '')}",
                        )
                        docs.append(doc)

                    # Pagination
                    messages_url = msg_data.get("@odata.nextLink")

                progress.update(
                    task,
                    completed=1,
                    total=1,
                    description=(
                        f"[cyan]Fetched {team_name}/{channel_name} "
                        f"━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━ 100%"
                    ),
                )

    console.print(
        f"Teams: loaded {len(docs)} messages from {len(team_ids)} team(s)"
    )

    return docs


def _write_teams_env(
    client_id: str,
    client_secret: str,
    tenant_id: str,
    team_ids: str,
) -> None:
    env_path = Path(".env")
    env_lines: list[str] = []
    if env_path.exists():
        env_lines = env_path.read_text(encoding="utf-8").splitlines()

    env_lines = [
        line
        for line in env_lines
        if not line.startswith(
            (
                "TEAMS_CLIENT_ID=",
                "TEAMS_CLIENT_SECRET=",
                "TEAMS_TENANT_ID=",
                "TEAMS_TEAM_IDS=",
            )
        )
    ]

    env_lines.append(f"TEAMS_CLIENT_ID={client_id}")
    env_lines.append(f"TEAMS_CLIENT_SECRET={client_secret}")
    env_lines.append(f"TEAMS_TENANT_ID={tenant_id}")
    env_lines.append(f"TEAMS_TEAM_IDS={team_ids}")

    env_path.write_text("\n".join(env_lines) + "\n", encoding="utf-8")
    print("[KORE] Saved Microsoft Teams credentials to .env")


def _load_github() -> list[Document]:
    from llama_index.readers.github import (
        GithubRepositoryReader,
        GitHubRepositoryIssuesReader,
    )
    from llama_index.readers.github.repository.github_client import (
        GithubClient,
    )
    from llama_index.readers.github.issues.github_client import (
        GitHubIssuesClient,
    )

    import click as _click
    from rich.console import Console
    from rich.panel import Panel
    from rich.progress import (
        Progress,
        SpinnerColumn,
        TextColumn,
    )

    console = Console()

    token = os.getenv("GITHUB_TOKEN")
    repo_spec = os.getenv("GITHUB_REPO")
    branch = os.getenv("GITHUB_BRANCH", "main")
    extensions_raw = os.getenv("GITHUB_EXTENSIONS")

    if not token or not repo_spec:
        console.print(Panel("[bold]GitHub Setup[/bold]", style="bold blue"))
        console.print("[dim]Required: Personal Access Token + repo to ingest.[/dim]\n")

        if not token:
            token = _click.prompt(
                "Enter your GitHub Personal Access Token",
                hide_input=True,
            )
        if not repo_spec:
            repo_spec = _click.prompt(
                "Enter repo to ingest (format: owner/repo)"
            )
        if not os.getenv("GITHUB_BRANCH"):
            branch = _click.prompt(
                "Enter branch", default="main", show_default=True
            )
        if not extensions_raw:
            extensions_raw = _click.prompt(
                "File extensions to include",
                default=".py .md .yaml .yml .ts .js",
                show_default=True,
            )

        _write_github_env(token, repo_spec, branch, extensions_raw)

    parts = repo_spec.split("/")
    if len(parts) != 2:
        raise ValueError(
            f"Invalid repo format '{repo_spec}'. Expected 'owner/repo'."
        )
    owner, repo_name = parts

    extensions = [
        ext.strip() if ext.strip().startswith(".") else f".{ext.strip()}"
        for ext in extensions_raw.split()
    ]

    repo_client = GithubClient(github_token=token)
    issues_client = GitHubIssuesClient(github_token=token)

    repo_reader = GithubRepositoryReader(
        github_client=repo_client,
        owner=owner,
        repo=repo_name,
        filter_directories=(
            [
                "node_modules",
                ".git",
                "__pycache__",
                "dist",
                "build",
                "venv",
                ".venv",
                "target",
                ".next",
                "coverage",
                ".tox",
                ".eggs",
                ".mypy_cache",
                ".pytest_cache",
                ".ruff_cache",
                "vendor",
                "bower_components",
                ".yarn",
                ".cache",
                ".terraform",
            ],
            GithubRepositoryReader.FilterType.EXCLUDE,
        ),
        filter_file_extensions=(
            extensions,
            GithubRepositoryReader.FilterType.INCLUDE,
        ),
        timeout=30,
        fail_on_error=False,
    )

    issues_reader = GitHubRepositoryIssuesReader(
        github_client=issues_client,
        owner=owner,
        repo=repo_name,
    )

    all_docs: list[Document] = []

    with Progress(
        SpinnerColumn(),
        TextColumn("[cyan]{task.description}"),
        console=console,
    ) as progress:
        task_repo = progress.add_task(
            f"Fetching GitHub repo files...",
            total=None,
        )
        try:
            repo_docs = repo_reader.load_data(branch=branch)
        except Exception as exc:
            console.print(f"[red]Error loading repo files: {exc}[/red]")
            repo_docs = []

        for doc in repo_docs:
            metadata = dict(doc.metadata)
            metadata["source_type"] = "github"
            metadata["repo"] = repo_spec
            metadata["branch"] = branch
            doc.metadata = metadata

        all_docs.extend(repo_docs)
        progress.update(
            task_repo,
            completed=1,
            total=1,
            description=(
                f"[cyan]Loaded {len(repo_docs)} files "
                f"━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━ 100%"
            ),
        )

        task_issues = progress.add_task(
            "Fetching GitHub issues...",
            total=None,
        )
        try:
            issue_docs = []
            for state in (
                GitHubRepositoryIssuesReader.IssueState.OPEN,
                GitHubRepositoryIssuesReader.IssueState.CLOSED,
            ):
                state_issues = issues_reader.load_data(state=state)
                issue_docs.extend(state_issues)
        except Exception:
            issue_docs = []

        for doc in issue_docs:
            metadata = dict(doc.metadata)
            metadata["source_type"] = "github"
            metadata["repo"] = repo_spec
            metadata["branch"] = branch
            metadata["issue_number"] = str(doc.doc_id) if doc.doc_id else ""
            doc.metadata = metadata

        all_docs.extend(issue_docs)
        progress.update(
            task_issues,
            completed=1,
            total=1,
            description=(
                f"[cyan]Loaded {len(issue_docs)} issues "
                f"━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━ 100%"
            ),
        )

    console.print(
        f"[bold]GitHub:[/bold] loaded {len(all_docs)} documents "
        f"from {repo_spec} ({branch} branch)"
    )

    return all_docs


def _write_github_env(
    token: str, repo: str, branch: str, extensions: str
) -> None:
    env_path = Path(".env")
    env_lines: list[str] = []
    if env_path.exists():
        env_lines = env_path.read_text(encoding="utf-8").splitlines()

    env_lines = [
        line
        for line in env_lines
        if not line.startswith(
            ("GITHUB_TOKEN=", "GITHUB_REPO=", "GITHUB_BRANCH=", "GITHUB_EXTENSIONS=")
        )
    ]

    env_lines.append(f"GITHUB_TOKEN={token}")
    env_lines.append(f"GITHUB_REPO={repo}")
    env_lines.append(f"GITHUB_BRANCH={branch}")
    env_lines.append(f"GITHUB_EXTENSIONS={extensions}")

    env_path.write_text("\n".join(env_lines) + "\n", encoding="utf-8")
    print("[KORE] Saved GitHub credentials to .env")


def _load_notion() -> list[Document]:
    from llama_index.readers.notion import NotionPageReader

    import click as _click
    from rich.console import Console
    from rich.panel import Panel
    from rich.progress import (
        Progress,
        SpinnerColumn,
        TextColumn,
    )

    console = Console()

    token = os.getenv("NOTION_TOKEN") or os.getenv("NOTION_INTEGRATION_TOKEN")

    if not token:
        console.print(Panel("[bold]Notion Setup[/bold]", style="bold blue"))
        console.print(
            "[dim]A Notion integration token is required to read pages.[/dim]\n"
        )
        console.print("""
[bold]How to create a Notion integration:[/bold]
1. Go to [link]https://www.notion.so/my-integrations[/link]
2. Click [bold]+ New integration[/bold]
3. Name it "KORE Ingestion" and select your workspace
4. Click [bold]Submit[/bold] and copy the [bold]Internal Integration Secret[/bold]
5. Go to each page/database you want to ingest → [bold]...[/bold] → [bold]Connections[/bold]
6. Add "KORE Ingestion" to grant access
""")
        token = _click.prompt(
            "Enter your Notion Integration Token",
            hide_input=True,
        )
        _write_notion_env(token)

    reader = NotionPageReader(integration_token=token)

    all_docs: list[Document] = []

    with Progress(
        SpinnerColumn(),
        TextColumn("[cyan]{task.description}"),
        console=console,
    ) as progress:
        task_pages = progress.add_task(
            "Fetching Notion pages...",
            total=None,
        )
        try:
            docs = reader.load_data(load_all_if_empty=True)
        except Exception as exc:
            console.print(
                f"[red]Error loading Notion pages: {exc}[/red]\n"
                f"[dim]Make sure the integration has access to the pages/databases "
                f"you want to ingest. In Notion, open each page → ... → Connections → "
                f"add your integration.[/dim]"
            )
            docs = []

        for doc in docs:
            metadata = dict(doc.metadata) if doc.metadata else {}
            existing_page_id = (
                metadata.get("page_id")
                or str(doc.doc_id)
            )
            metadata["source_type"] = "notion"
            metadata["page_id"] = existing_page_id
            metadata["page_title"] = ""
            doc.metadata = metadata

        all_docs.extend(docs)
        progress.update(
            task_pages,
            completed=1,
            total=1,
            description=(
                f"[cyan]Loaded {len(docs)} pages "
                f"━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━ 100%"
            ),
        )

    console.print(
        f"[bold]Notion:[/bold] loaded {len(all_docs)} pages"
    )

    return all_docs


def _write_notion_env(token: str) -> None:
    env_path = Path(".env")
    env_lines: list[str] = []
    if env_path.exists():
        env_lines = env_path.read_text(encoding="utf-8").splitlines()

    env_lines = [
        line
        for line in env_lines
        if not line.startswith(("NOTION_TOKEN=", "NOTION_INTEGRATION_TOKEN="))
    ]

    env_lines.append(f"NOTION_TOKEN={token}")

    env_path.write_text("\n".join(env_lines) + "\n", encoding="utf-8")
    print("[KORE] Saved Notion credentials to .env")


def _load_jira() -> list[Document]:
    from llama_index.readers.jira import JiraReader

    import click as _click
    from rich.console import Console
    from rich.panel import Panel
    from rich.progress import (
        Progress,
        SpinnerColumn,
        TextColumn,
    )

    console = Console()

    server_url = os.getenv("JIRA_URL", "")
    username = os.getenv("JIRA_USERNAME", "")
    api_token = os.getenv("JIRA_API_TOKEN", "")
    project_keys_raw = os.getenv("JIRA_PROJECT_KEYS", "")

    if not all((server_url, username, api_token, project_keys_raw)):
        console.print(Panel("[bold]Jira Setup[/bold]", style="bold blue"))
        console.print(
            "[dim]Jira Cloud credentials are required to read issues.[/dim]\n"
        )
        console.print("""
[bold]How to get your Jira Cloud credentials:[/bold]

[bold]1. Server URL[/bold]
   Open Jira in your browser. The URL is something like:
   [cyan]https://yourcompany.atlassian.net[/cyan]
   Copy that exact URL.

[bold]2. Username[/bold]
   Use the email address you log into Atlassian with.

[bold]3. API Token[/bold]
   a. Go to [link]https://id.atlassian.com/manage-profile/security/api-tokens[/link]
   b. Click [bold]Create API token[/bold]
   c. Label: "KORE Ingestion"
   d. Click [bold]Create[/bold] — a modal appears with your token
   e. Click [bold]Copy[/bold] immediately (you cannot view it again)
   f. The token is ~70 characters long, starts with letters/numbers

[bold]4. Project Key(s)[/bold]
   In Jira, open any project. The key is the short prefix on tickets
   (e.g. [bold]KAN-123[/bold] → project key is [bold]KAN[/bold]).
   You can find it in Project Settings → Details.
""")
        if not server_url:
            server_url = _click.prompt(
                "Jira server URL (e.g. https://company.atlassian.net)"
            )
        if not username:
            username = _click.prompt("Jira username (your Atlassian email)")
        if not api_token:
            api_token = _click.prompt(
                "Jira API token (the ~70 char string from the API token page)",
                hide_input=True,
            )
        if not project_keys_raw:
            project_keys_raw = _click.prompt(
                "Project keys (comma-separated, e.g. KAN,PROJ)"
            )

        _write_jira_env(server_url, username, api_token, project_keys_raw)

    # Strip protocol if present — JiraReader prepends https:// internally
    domain = server_url.replace("https://", "").replace("http://", "").strip("/")

    project_keys = [k.strip() for k in project_keys_raw.split(",") if k.strip()]
    if len(project_keys) == 1:
        jql = f"project = {project_keys[0]}"
    else:
        jql = f"project in ({', '.join(project_keys)})"

    reader = JiraReader(
        email=username,
        api_token=api_token,
        server_url=domain,
    )

    all_docs: list[Document] = []

    with Progress(
        SpinnerColumn(),
        TextColumn("[cyan]{task.description}"),
        console=console,
    ) as progress:
        task_jira = progress.add_task(
            f"Fetching Jira issues ({', '.join(project_keys)})...",
            total=None,
        )

        try:
            # NOTE: the jira library auto-paginates internally.
            # max_results=50 here is the *page size* for each API call,
            # not a hard limit on total issues returned.
            docs = reader.load_data(query=jql, start_at=0, max_results=50)

            for doc in docs:
                metadata = dict(doc.metadata) if doc.metadata else {}
                metadata["source_type"] = "jira"
                metadata["ticket_id"] = str(doc.doc_id) if doc.doc_id else ""
                metadata["status"] = metadata.get("status", "")
                metadata["assignee"] = metadata.get("assignee", "")
                doc.metadata = metadata

            all_docs.extend(docs)

        except Exception as exc:
            console.print(f"[red]Error loading Jira issues: {exc}[/red]")

        progress.update(
            task_jira,
            completed=1,
            total=1,
            description=(
                f"[cyan]Loaded {len(all_docs)} issues "
                f"━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━ 100%"
            ),
        )

    console.print(
        f"[bold]Jira:[/bold] loaded {len(all_docs)} issues "
        f"from projects: {', '.join(project_keys)}"
    )

    return all_docs


def _write_jira_env(
    server_url: str, username: str, api_token: str, project_keys: str
) -> None:
    env_path = Path(".env")
    env_lines: list[str] = []
    if env_path.exists():
        env_lines = env_path.read_text(encoding="utf-8").splitlines()

    env_lines = [
        line
        for line in env_lines
        if not line.startswith(
            (
                "JIRA_URL=",
                "JIRA_USERNAME=",
                "JIRA_API_TOKEN=",
                "JIRA_PROJECT_KEYS=",
            )
        )
    ]

    env_lines.append(f"JIRA_URL={server_url}")
    env_lines.append(f"JIRA_USERNAME={username}")
    env_lines.append(f"JIRA_API_TOKEN={api_token}")
    env_lines.append(f"JIRA_PROJECT_KEYS={project_keys}")

    env_path.write_text("\n".join(env_lines) + "\n", encoding="utf-8")
    print("[KORE] Saved Jira credentials to .env")


def _load_confluence() -> list[Document]:
    from llama_index.readers.confluence import ConfluenceReader

    import click as _click
    from rich.console import Console
    from rich.panel import Panel
    from rich.progress import (
        Progress,
        SpinnerColumn,
        TextColumn,
    )

    console = Console()

    url = os.getenv("CONFLUENCE_URL", "")
    username = os.getenv("CONFLUENCE_USERNAME", "")
    api_token = os.getenv("CONFLUENCE_API_TOKEN", "")
    space_keys_raw = os.getenv("CONFLUENCE_SPACE_KEYS", "")

    if not all((url, username, api_token, space_keys_raw)):
        console.print(Panel("[bold]Confluence Setup[/bold]", style="bold blue"))
        console.print(
            "[dim]Confluence Cloud credentials are required to read pages.[/dim]\n"
        )
        console.print("""
[bold]How to get your Confluence credentials:[/bold]

[bold]1. Base URL[/bold]
   Open Confluence in your browser. The URL is something like:
   [cyan]https://yourcompany.atlassian.net/wiki[/cyan]
   Copy that exact URL (must end with /wiki).

[bold]2. Username[/bold]
   Use the email address you log into Atlassian with.

[bold]3. API Token[/bold]
   a. Go to [link]https://id.atlassian.com/manage-profile/security/api-tokens[/link]
   b. Click [bold]Create API token[/bold]
   c. Label: "KORE Ingestion"
   d. Click [bold]Create[/bold] -- a modal appears with your token
   e. Click [bold]Copy[/bold] immediately (you cannot view it again)

[bold]4. Space Key(s)[/bold]
   In Confluence, open any space. The key is the short prefix in the URL
   (e.g. [bold]https://.../spaces/PROJ[/bold] → space key is [bold]PROJ[/bold]).
""")
        if not url:
            url = _click.prompt(
                "Confluence base URL (e.g. https://company.atlassian.net/wiki)"
            )
        if not username:
            username = _click.prompt("Confluence username (your Atlassian email)")
        if not api_token:
            api_token = _click.prompt(
                "Confluence API token (the ~70 char string from the API token page)",
                hide_input=True,
            )
        if not space_keys_raw:
            space_keys_raw = _click.prompt(
                "Space keys (comma-separated, e.g. PROJ,ENG)"
            )

        _write_confluence_env(url, username, api_token, space_keys_raw)

    space_keys = [k.strip() for k in space_keys_raw.split(",") if k.strip()]

    reader = ConfluenceReader(
        base_url=url,
        api_token=api_token,
    )

    all_docs: list[Document] = []

    with Progress(
        SpinnerColumn(),
        TextColumn("[cyan]{task.description}"),
        console=console,
    ) as progress:
        for space_key in space_keys:
            task = progress.add_task(
                f"Fetching Confluence space {space_key}...",
                total=None,
            )
            try:
                docs = reader.load_data(
                    space_key=space_key,
                    page_status="current",
                    include_attachments=False,
                )
            except Exception as exc:
                console.print(f"[red]Error loading Confluence space {space_key}: {exc}[/red]")
                docs = []

            for doc in docs:
                metadata = dict(doc.metadata) if doc.metadata else {}
                page_id = str(doc.doc_id) if doc.doc_id else ""
                title = metadata.get("title", "")
                space = metadata.get("space_key", space_key)
                modified = metadata.get("last_modified", "")
                metadata["source_type"] = "confluence"
                metadata["space_key"] = space
                metadata["page_title"] = title
                metadata["page_id"] = page_id
                metadata["last_modified"] = modified
                doc.metadata = metadata

            all_docs.extend(docs)
            progress.update(
                task,
                completed=1,
                total=1,
                description=(
                    f"[cyan]Loaded {len(docs)} pages from {space_key} "
                    f"━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━ 100%"
                ),
            )

    console.print(
        f"Confluence: loaded {len(all_docs)} pages from spaces: {', '.join(space_keys)}"
    )

    return all_docs


def _write_confluence_env(
    url: str, username: str, api_token: str, space_keys: str
) -> None:
    env_path = Path(".env")
    env_lines: list[str] = []
    if env_path.exists():
        env_lines = env_path.read_text(encoding="utf-8").splitlines()

    env_lines = [
        line
        for line in env_lines
        if not line.startswith(
            (
                "CONFLUENCE_URL=",
                "CONFLUENCE_USERNAME=",
                "CONFLUENCE_API_TOKEN=",
                "CONFLUENCE_SPACE_KEYS=",
            )
        )
    ]

    env_lines.append(f"CONFLUENCE_URL={url}")
    env_lines.append(f"CONFLUENCE_USERNAME={username}")
    env_lines.append(f"CONFLUENCE_API_TOKEN={api_token}")
    env_lines.append(f"CONFLUENCE_SPACE_KEYS={space_keys}")

    env_path.write_text("\n".join(env_lines) + "\n", encoding="utf-8")
    print("[KORE] Saved Confluence credentials to .env")


def _load_linear() -> list[Document]:
    """Load issues, comments, project updates, and documents from Linear via GraphQL API.

    Requires environment variables:
      LINEAR_API_KEY     - Personal API key from Linear
      LINEAR_TEAM_NAMES  - Comma-separated list of team names (optional, empty = all)

    Create an API key at: linear.app → Settings → API → Personal API keys
    """
    import requests

    import click as _click
    from rich.console import Console
    from rich.panel import Panel
    from rich.progress import (
        Progress,
        SpinnerColumn,
        TextColumn,
    )

    console = Console()

    api_key = os.getenv("LINEAR_API_KEY")
    team_names_raw = os.getenv("LINEAR_TEAM_NAMES", "")

    if not api_key:
        console.print(Panel("[bold]Linear Setup[/bold]", style="bold blue"))
        print("Go to linear.app → Settings → API → Personal API keys")
        api_key = _click.prompt("Enter your Linear API key:", hide_input=True)
        team_names_raw = _click.prompt(
            "Enter team names to ingest (comma-separated, or Enter for all):",
            default="",
            show_default=False,
        )
        _write_linear_env(api_key, team_names_raw)

    team_names = [t.strip() for t in team_names_raw.split(",") if t.strip()]

    headers = {
        "Authorization": api_key,
        "Content-Type": "application/json",
    }

    def _graphql(query: str, variables: dict | None = None) -> dict:
        resp = requests.post(
            "https://api.linear.app/graphql",
            json={"query": query, "variables": variables or {}},
            headers=headers,
            timeout=30,
        )
        resp.raise_for_status()
        data = resp.json()
        if "errors" in data:
            raise RuntimeError(f"Linear GraphQL error: {data['errors']}")
        return data.get("data", {})

    # Fetch teams to map names to IDs
    teams_data = _graphql("""
        query {
            teams {
                nodes {
                    id
                    name
                    key
                }
            }
        }
    """)
    all_teams = {t["name"].lower(): t for t in teams_data.get("teams", {}).get("nodes", [])}

    target_teams = []
    if team_names:
        for name in team_names:
            team = all_teams.get(name.lower())
            if team:
                target_teams.append(team)
    else:
        target_teams = list(all_teams.values())

    docs: list[Document] = []

    with Progress(
        SpinnerColumn(),
        TextColumn("[cyan]{task.description}"),
        console=console,
    ) as progress:
        for team in target_teams:
            team_id = team["id"]
            team_name = team["name"]
            team_key = team.get("key", team_name)

            task = progress.add_task(
                f"Fetching Linear team {team_name}...",
                total=None,
            )

            # Fetch issues for this team
            issues_query = """
                query($teamId: ID!) {
                    issues(filter: { team: { id: { eq: $teamId } } }) {
                        nodes {
                            id
                            identifier
                            title
                            description
                            state { name }
                            assignee { name email }
                            comments { nodes { bodyText user { name } } }
                            project { name }
                            createdAt
                            updatedAt
                        }
                    }
                }
            """
            try:
                issues_data = _graphql(issues_query, {"teamId": team_id})
            except Exception as exc:
                console.print(f"[red]Error fetching issues for {team_name}: {exc}[/red]")
                issues_data = {}

            issues = issues_data.get("issues", {}).get("nodes", [])

            for issue in issues:
                issue_id = issue.get("identifier", issue.get("id", ""))
                title = issue.get("title", "")
                description = issue.get("description", "") or ""
                status = issue.get("state", {}).get("name", "")
                assignee = issue.get("assignee", {}) or {}
                assignee_name = assignee.get("name", "") if assignee else ""
                project_name = issue.get("project", {}).get("name", "") if issue.get("project") else ""
                created = issue.get("createdAt", "")
                updated = issue.get("updatedAt", "")

                # Build comments list
                comments = issue.get("comments", {}).get("nodes", [])
                comment_texts = []
                for c in comments:
                    body = c.get("bodyText", "")
                    author = c.get("user", {}).get("name", "") if c.get("user") else ""
                    if body:
                        comment_texts.append({"text": body, "author": author})

                # Main document for the issue (title + description)
                issue_text = f"{title}\n\n{description}".strip()
                if issue_text:
                    metadata = {
                        "source_type": "linear",
                        "team_name": team_name,
                        "team_key": team_key,
                        "issue_id": issue_id,
                        "issue_title": title,
                        "project_name": project_name,
                        "status": status,
                        "assignee": assignee_name,
                        "created_at": created,
                        "last_modified": updated,
                        "comments": comment_texts,
                    }
                    doc = Document(
                        text=issue_text,
                        metadata=metadata,
                        id_=f"linear_{team_key}_{issue_id}",
                    )
                    docs.append(doc)

            # Fetch project updates for this team
            projects_query = """
                query($teamId: ID!) {
                    projects(filter: { teams: { id: { eq: $teamId } } }) {
                        nodes {
                            id
                            name
                            description
                            state
                            updatedAt
                        }
                    }
                }
            """
            try:
                projects_data = _graphql(projects_query, {"teamId": team_id})
            except Exception as exc:
                console.print(f"[red]Error fetching projects for {team_name}: {exc}[/red]")
                projects_data = {}

            projects = projects_data.get("projects", {}).get("nodes", [])
            for project in projects:
                proj_name = project.get("name", "")
                proj_desc = project.get("description", "") or ""
                proj_state = project.get("state", "")
                proj_updated = project.get("updatedAt", "")
                if proj_desc:
                    metadata = {
                        "source_type": "linear",
                        "team_name": team_name,
                        "team_key": team_key,
                        "project_name": proj_name,
                        "project_id": project.get("id", ""),
                        "status": proj_state,
                        "assignee": "",
                        "last_modified": proj_updated,
                    }
                    doc = Document(
                        text=proj_desc,
                        metadata=metadata,
                        id_=f"linear_project_{team_key}_{project.get('id', '')}",
                    )
                    docs.append(doc)

            # Fetch document pages for this team
            docs_query = """
                query($teamId: ID!) {
                    documents(filter: { team: { id: { eq: $teamId } } }) {
                        nodes {
                            id
                            title
                            content
                            updatedAt
                        }
                    }
                }
            """
            try:
                documents_data = _graphql(docs_query, {"teamId": team_id})
            except Exception as exc:
                console.print(f"[red]Error fetching documents for {team_name}: {exc}[/red]")
                documents_data = {}

            document_nodes = documents_data.get("documents", {}).get("nodes", [])
            for doc_node in document_nodes:
                doc_title = doc_node.get("title", "")
                doc_content = doc_node.get("content", "") or ""
                doc_updated = doc_node.get("updatedAt", "")
                if doc_content:
                    metadata = {
                        "source_type": "linear",
                        "team_name": team_name,
                        "team_key": team_key,
                        "project_name": doc_title,
                        "document_id": doc_node.get("id", ""),
                        "status": "",
                        "assignee": "",
                        "last_modified": doc_updated,
                    }
                    doc = Document(
                        text=doc_content,
                        metadata=metadata,
                        id_=f"linear_doc_{team_key}_{doc_node.get('id', '')}",
                    )
                    docs.append(doc)

            progress.update(
                task,
                completed=1,
                total=1,
                description=(
                    f"[cyan]Loaded {len(issues)} issues from {team_name} "
                    f"━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━ 100%"
                ),
            )

    console.print(
        f"Linear: loaded {len(docs)} issues from {len(target_teams)} team(s)"
    )

    return docs


def _write_linear_env(api_key: str, team_names: str) -> None:
    env_path = Path(".env")
    env_lines: list[str] = []
    if env_path.exists():
        env_lines = env_path.read_text(encoding="utf-8").splitlines()

    env_lines = [
        line
        for line in env_lines
        if not line.startswith(
            (
                "LINEAR_API_KEY=",
                "LINEAR_TEAM_NAMES=",
            )
        )
    ]

    env_lines.append(f"LINEAR_API_KEY={api_key}")
    env_lines.append(f"LINEAR_TEAM_NAMES={team_names}")

    env_path.write_text("\n".join(env_lines) + "\n", encoding="utf-8")
    print("[KORE] Saved Linear credentials to .env")


def _load_zendesk() -> list[Document]:
    """Load macros, tickets (with comments), and help center articles from Zendesk.

    Requires environment variables:
      ZENDESK_SUBDOMAIN - e.g. "mycompany"
      ZENDESK_EMAIL     - Agent email address
      ZENDESK_API_TOKEN - API token from Admin → Channels → API

    Uses ZendeskReader (LlamaHub) for help center articles and the Zendesk REST
    API directly for tickets, comments, and macros.
    """
    import requests
    from llama_index.readers.zendesk import ZendeskReader

    import click as _click
    from rich.console import Console
    from rich.panel import Panel
    from rich.progress import (
        Progress,
        SpinnerColumn,
        TextColumn,
    )

    console = Console()

    subdomain = os.getenv("ZENDESK_SUBDOMAIN", "")
    email = os.getenv("ZENDESK_EMAIL", "")
    api_token = os.getenv("ZENDESK_API_TOKEN", "")

    if not all((subdomain, email, api_token)):
        console.print(Panel("[bold]Zendesk Setup[/bold]", style="bold blue"))
        console.print(
            "[dim]Zendesk credentials are required to read tickets, macros, and articles.[/dim]\n"
        )
        if not subdomain:
            subdomain = _click.prompt(
                "Zendesk subdomain (e.g. mycompany for https://mycompany.zendesk.com)"
            )
        if not email:
            email = _click.prompt("Zendesk agent email")
        if not api_token:
            api_token = _click.prompt(
                "Zendesk API token (found at Admin → Channels → API)",
                hide_input=True,
            )
        _write_zendesk_env(subdomain, email, api_token)

    base_url = f"https://{subdomain}.zendesk.com"
    auth = (f"{email}/token", api_token)

    def _zendesk_get(path: str, params: dict | None = None) -> dict:
        resp = requests.get(
            f"{base_url}/api/v2{path}",
            auth=auth,
            params=params,
            timeout=30,
        )
        resp.raise_for_status()
        return resp.json()

    docs: list[Document] = []
    macro_count = 0
    ticket_count = 0
    article_count = 0

    with Progress(
        SpinnerColumn(),
        TextColumn("[cyan]{task.description}"),
        console=console,
    ) as progress:

        # ---------- MACROS (highest priority — contain policy rules) ----------
        task_macros = progress.add_task("Fetching Zendesk macros...", total=None)
        try:
            macros_data = _zendesk_get("/macros.json")
            macros = macros_data.get("macros", [])
            for macro in macros:
                macro_id = str(macro.get("id", ""))
                title = macro.get("title", "")
                actions = macro.get("actions", [])
                action_texts = []
                for action in actions:
                    field = action.get("field", "")
                    value = action.get("value", "")
                    action_texts.append(f"{field}: {value}")
                text = f"Macro: {title}\n" + "\n".join(action_texts)
                if text.strip():
                    metadata = {
                        "source_type": "zendesk",
                        "macro_id": macro_id,
                        "doc_type": "macro",
                        "title": title,
                        "status": "active",
                    }
                    docs.append(
                        Document(
                            text=text,
                            metadata=metadata,
                            id_=f"zendesk_macro_{macro_id}",
                        )
                    )
                    macro_count += 1
        except Exception as exc:
            console.print(f"[red]Error loading Zendesk macros: {exc}[/red]")
        progress.update(
            task_macros,
            completed=1,
            total=1,
            description=f"[cyan]Loaded {macro_count} macros ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━ 100%",
        )

        # ---------- TICKETS (with comments — contain agent decisions) ----------
        task_tickets = progress.add_task("Fetching Zendesk tickets...", total=None)
        try:
            # Fetch recent tickets (last 90 days by default)
            from datetime import timedelta
            import time

            since_dt = datetime.now(timezone.utc) - timedelta(days=90)
            since_unix = int(since_dt.timestamp())

            tickets_data = _zendesk_get(
                "/tickets.json",
                params={"query": f"updated>{since_unix}", "sort_by": "updated_at", "sort_order": "desc"},
            )
            tickets = tickets_data.get("tickets", [])

            for ticket in tickets:
                ticket_id = str(ticket.get("id", ""))
                subject = ticket.get("subject", "")
                description = ticket.get("description", "")
                status = ticket.get("status", "")
                assignee_id = str(ticket.get("assignee_id", ""))
                updated = ticket.get("updated_at", "")

                # Fetch comments for this ticket
                comments_data = _zendesk_get(f"/tickets/{ticket_id}/comments.json")
                comments = comments_data.get("comments", [])
                comment_texts = []
                for c in comments:
                    body = c.get("body", "")
                    author_id = str(c.get("author_id", ""))
                    if body:
                        comment_texts.append({"text": body, "author_id": author_id})

                # Build ticket document (subject + description + comments list in metadata)
                ticket_text = f"Ticket {ticket_id}: {subject}\n\n{description}".strip()
                if ticket_text:
                    metadata = {
                        "source_type": "zendesk",
                        "ticket_id": ticket_id,
                        "doc_type": "ticket",
                        "title": subject,
                        "status": status,
                        "assignee_id": assignee_id,
                        "last_modified": updated,
                        "comments": comment_texts,
                    }
                    docs.append(
                        Document(
                            text=ticket_text,
                            metadata=metadata,
                            id_=f"zendesk_ticket_{ticket_id}",
                        )
                    )
                    ticket_count += 1
        except Exception as exc:
            console.print(f"[red]Error loading Zendesk tickets: {exc}[/red]")
        progress.update(
            task_tickets,
            completed=1,
            total=1,
            description=f"[cyan]Loaded {ticket_count} tickets ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━ 100%",
        )

        # ---------- HELP CENTER ARTICLES ----------
        task_articles = progress.add_task("Fetching Zendesk articles...", total=None)
        try:
            reader = ZendeskReader(
                zendesk_subdomain=subdomain,
                locale="en-us",
            )
            article_docs = reader.load_data()
            for doc in article_docs:
                metadata = dict(doc.metadata) if doc.metadata else {}
                article_id = str(doc.doc_id) if doc.doc_id else ""
                title = metadata.get("title", "")
                metadata["source_type"] = "zendesk"
                metadata["article_id"] = article_id
                metadata["doc_type"] = "article"
                metadata["title"] = title
                metadata["status"] = "published"
                doc.metadata = metadata
                docs.append(doc)
                article_count += 1
        except Exception as exc:
            console.print(f"[red]Error loading Zendesk articles: {exc}[/red]")
        progress.update(
            task_articles,
            completed=1,
            total=1,
            description=f"[cyan]Loaded {article_count} articles ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━ 100%",
        )

    console.print(
        f"Zendesk: loaded {len(docs)} documents "
        f"({macro_count} macros, {ticket_count} tickets, {article_count} articles)"
    )

    return docs


def _write_zendesk_env(
    subdomain: str, email: str, api_token: str
) -> None:
    env_path = Path(".env")
    env_lines: list[str] = []
    if env_path.exists():
        env_lines = env_path.read_text(encoding="utf-8").splitlines()

    env_lines = [
        line
        for line in env_lines
        if not line.startswith(
            (
                "ZENDESK_SUBDOMAIN=",
                "ZENDESK_EMAIL=",
                "ZENDESK_API_TOKEN=",
            )
        )
    ]

    env_lines.append(f"ZENDESK_SUBDOMAIN={subdomain}")
    env_lines.append(f"ZENDESK_EMAIL={email}")
    env_lines.append(f"ZENDESK_API_TOKEN={api_token}")

    env_path.write_text("\n".join(env_lines) + "\n", encoding="utf-8")
    print("[KORE] Saved Zendesk credentials to .env")


def _load_google_drive() -> list[Document]:
    """Load files from Google Drive using GoogleDriveReader from LlamaHub.

    Requires environment variables:
      GOOGLE_DRIVE_CREDENTIALS_FILE - Path to GCP service account JSON key
      GOOGLE_DRIVE_FOLDER_IDS        - Comma-separated list of Drive folder IDs

    The service account must have Drive API enabled and folders shared with it.
    """
    from llama_index.readers.google import GoogleDriveReader

    import click as _click
    from rich.console import Console
    from rich.panel import Panel
    from rich.progress import (
        Progress,
        SpinnerColumn,
        TextColumn,
    )

    console = Console()

    creds_path_str = os.getenv("GOOGLE_DRIVE_CREDENTIALS_FILE", "")
    folder_ids_raw = os.getenv("GOOGLE_DRIVE_FOLDER_IDS", "")

    if not all((creds_path_str, folder_ids_raw)):
        console.print(Panel("[bold]Google Drive Setup[/bold]", style="bold blue"))
        console.print(
            "[dim]Google Drive credentials are required to read files.[/dim]\n"
        )
        console.print("""
[bold]How to set up Google Drive access:[/bold]

[bold]1. Create a GCP project[/bold]
   Go to [link]https://console.cloud.google.com/[/link] and create a new project.

[bold]2. Enable the Drive API[/bold]
   APIs & Services -> Library -> Search "Google Drive API" -> Enable.

[bold]3. Create a service account[/bold]
   APIs & Services -> Credentials -> Create credentials -> Service account.
   Copy the service account email (looks like: ...@...gserviceaccount.com).

[bold]4. Create a JSON key[/bold]
   Keys -> Add Key -> Create new key -> JSON.
   Download the .json file to this machine.

[bold]5. Share your Drive folders[/bold]
   In Google Drive, right-click each folder you want to ingest -> Share.
   Paste the service account email and grant Viewer access.
""")
        if not creds_path_str:
            creds_path_str = _click.prompt(
                "Enter path to your service account credentials JSON file:"
            )
        if not folder_ids_raw:
            folder_ids_raw = _click.prompt(
                "Enter Google Drive folder IDs to ingest (comma-separated):"
            )

        _write_google_drive_env(creds_path_str, folder_ids_raw)

    folder_ids = [f.strip() for f in folder_ids_raw.split(",") if f.strip()]

    # Read service account JSON
    creds_path = Path(creds_path_str)
    if not creds_path.exists():
        raise FileNotFoundError(f"Google Drive credentials file not found: {creds_path}")
    service_account_key = creds_path.read_text(encoding="utf-8")

    all_docs: list[Document] = []

    with Progress(
        SpinnerColumn(),
        TextColumn("[cyan]{task.description}"),
        console=console,
    ) as progress:
        for folder_id in folder_ids:
            task = progress.add_task(
                f"Fetching Google Drive folder {folder_id}...",
                total=None,
            )

            try:
                reader = GoogleDriveReader(
                    folder_id=folder_id,
                    service_account_key=service_account_key,
                )
                docs = reader.load_data()
            except Exception as exc:
                console.print(f"[red]Error loading Google Drive folder {folder_id}: {exc}[/red]")
                docs = []

            for doc in docs:
                metadata = dict(doc.metadata) if doc.metadata else {}
                file_id = str(doc.doc_id) if doc.doc_id else ""
                file_name = metadata.get("file_name", "")
                mime_type = metadata.get("mime_type", "")
                modified = metadata.get("modified_time", "")
                metadata["source_type"] = "google_drive"
                metadata["file_id"] = file_id
                metadata["file_name"] = file_name
                metadata["mime_type"] = mime_type
                metadata["last_modified"] = modified
                doc.metadata = metadata

            all_docs.extend(docs)
            progress.update(
                task,
                completed=1,
                total=1,
                description=(
                    f"[cyan]Loaded {len(docs)} files from {folder_id} "
                    f"━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━ 100%"
                ),
            )

    console.print(
        f"Google Drive: loaded {len(all_docs)} files from {len(folder_ids)} folder(s)"
    )

    return all_docs


def _write_google_drive_env(
    creds_path: str, folder_ids: str
) -> None:
    env_path = Path(".env")
    env_lines: list[str] = []
    if env_path.exists():
        env_lines = env_path.read_text(encoding="utf-8").splitlines()

    env_lines = [
        line
        for line in env_lines
        if not line.startswith(
            (
                "GOOGLE_DRIVE_CREDENTIALS_FILE=",
                "GOOGLE_DRIVE_FOLDER_IDS=",
            )
        )
    ]

    env_lines.append(f"GOOGLE_DRIVE_CREDENTIALS_FILE={creds_path}")
    env_lines.append(f"GOOGLE_DRIVE_FOLDER_IDS={folder_ids}")

    env_path.write_text("\n".join(env_lines) + "\n", encoding="utf-8")
    print("[KORE] Saved Google Drive credentials to .env")

def _load_documents(path: str) -> list[Document]:
    """Load local documents (PDF, Word, Markdown, text, CSV, JSON, YAML).

    Parameters
    ----------
    path : str
        Path to a single file or a directory. Directories are scanned recursively.

    Returns
    -------
    list[Document]
        Documents with metadata["source_type"] = "document" and file_type set.
    """
    import csv

    from rich.console import Console

    console = Console()

    root = Path(path)
    if not root.exists():
        raise FileNotFoundError(f"Path not found: {root}")

    # Supported extensions mapped to file_type
    EXT_MAP: dict[str, str] = {
        ".pdf": "pdf",
        ".docx": "docx",
        ".doc": "docx",
        ".md": "markdown",
        ".mdx": "markdown",
        ".txt": "text",
        ".csv": "csv",
        ".json": "json",
        ".yaml": "yaml",
        ".yml": "yaml",
    }

    EXCLUDED_DIRS = {
        "node_modules",
        ".git",
        "__pycache__",
        "venv",
        ".venv",
        "dist",
        "build",
        ".pytest_cache",
        ".mypy_cache",
        ".eggs",
        ".tox",
        ".ruff_cache",
        ".cache",
        ".next",
        "coverage",
        "target",
        ".terraform",
        ".yarn",
        "bower_components",
    }

    POLICY_KEYWORDS = [
        "policy", "procedure", "handbook", "guidelines",
        "rules", "standards", "compliance", "sop", "protocol",
    ]

    def _is_likely_policy(filename: str) -> bool:
        lowered = filename.lower()
        return any(kw in lowered for kw in POLICY_KEYWORDS)

    # Collect files
    if root.is_file():
        if root.suffix.lower() in EXT_MAP:
            files = [root]
        else:
            console.print(f"  [yellow]⚠ {root.name} (unsupported — skipped)[/yellow]")
            files = []
    else:
        files = []
        for item in root.rglob("*"):
            if not item.is_file():
                continue
            # Skip excluded directories
            if any(part in EXCLUDED_DIRS for part in item.parts):
                continue
            if item.suffix.lower() in EXT_MAP:
                files.append(item)

    console.print(f"Loading documents from {path}")

    docs: list[Document] = []
    for file_path in sorted(files, key=lambda p: p.name):
        ext = file_path.suffix.lower()
        file_type = EXT_MAP[ext]
        file_name = file_path.name
        file_size_kb = round(file_path.stat().st_size / 1024, 2)
        last_modified = datetime.fromtimestamp(
            file_path.stat().st_mtime, tz=timezone.utc
        ).isoformat()

        try:
            if file_type == "pdf":
                text, page_count = _extract_pdf(file_path)
                if not text.strip():
                    console.print(f"  [yellow]⚠ {file_name} (image-only PDF — skipped)[/yellow]")
                    continue
                metadata = {
                    "source_type": "document",
                    "file_type": "pdf",
                    "file_name": file_name,
                    "file_path": str(file_path),
                    "file_size_kb": file_size_kb,
                    "last_modified": last_modified,
                    "page_count": page_count,
                    "likely_explicit_policy": _is_likely_policy(file_name),
                }
                doc = Document(text=text, metadata=metadata, id_=f"doc_pdf_{file_path.stem}")
                docs.append(doc)
                console.print(f"  [green]✓ {file_name} ({page_count} pages)[/green]")

            elif file_type == "docx":
                text = _extract_docx(file_path)
                metadata = {
                    "source_type": "document",
                    "file_type": "docx",
                    "file_name": file_name,
                    "file_path": str(file_path),
                    "file_size_kb": file_size_kb,
                    "last_modified": last_modified,
                    "likely_explicit_policy": _is_likely_policy(file_name),
                }
                doc = Document(text=text, metadata=metadata, id_=f"doc_docx_{file_path.stem}")
                docs.append(doc)
                console.print(f"  [green]✓ {file_name}[/green]")

            elif file_type in ("markdown", "text"):
                text = file_path.read_text(encoding="utf-8")
                metadata = {
                    "source_type": "document",
                    "file_type": file_type,
                    "file_name": file_name,
                    "file_path": str(file_path),
                    "file_size_kb": file_size_kb,
                    "last_modified": last_modified,
                    "likely_explicit_policy": _is_likely_policy(file_name),
                }
                doc = Document(text=text, metadata=metadata, id_=f"doc_{file_type}_{file_path.stem}")
                docs.append(doc)
                console.print(f"  [green]✓ {file_name}[/green]")

            elif file_type == "csv":
                text = _extract_csv(file_path)
                metadata = {
                    "source_type": "document",
                    "file_type": "csv",
                    "file_name": file_name,
                    "file_path": str(file_path),
                    "file_size_kb": file_size_kb,
                    "last_modified": last_modified,
                    "likely_explicit_policy": _is_likely_policy(file_name),
                }
                doc = Document(text=text, metadata=metadata, id_=f"doc_csv_{file_path.stem}")
                docs.append(doc)
                console.print(f"  [green]✓ {file_name}[/green]")

            elif file_type in ("json", "yaml"):
                text = file_path.read_text(encoding="utf-8")
                metadata = {
                    "source_type": "document",
                    "file_type": file_type,
                    "file_name": file_name,
                    "file_path": str(file_path),
                    "file_size_kb": file_size_kb,
                    "last_modified": last_modified,
                    "likely_explicit_policy": _is_likely_policy(file_name),
                }
                doc = Document(text=text, metadata=metadata, id_=f"doc_{file_type}_{file_path.stem}")
                docs.append(doc)
                console.print(f"  [green]✓ {file_name}[/green]")

            else:
                console.print(f"  [yellow]⚠ {file_name} (unsupported — skipped)[/yellow]")

        except Exception as exc:
            console.print(f"  [red]✗ {file_name} ({exc})[/red]")

    console.print(f"Documents: loaded {len(docs)} files")
    return docs


def _extract_pdf(file_path: Path) -> tuple[str, int]:
    from pypdf import PdfReader

    reader = PdfReader(str(file_path))
    page_count = len(reader.pages)
    texts: list[str] = []
    for i, page in enumerate(reader.pages):
        page_text = page.extract_text() or ""
        if page_text.strip():
            texts.append(f"--- Page {i + 1} ---\n{page_text}")
    return "\n\n".join(texts), page_count


def _extract_docx(file_path: Path) -> str:
    from docx import Document as DocxDocument

    doc = DocxDocument(str(file_path))
    paragraphs: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Extract table content
    for table in doc.tables:
        table_texts: list[str] = []
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text:
                table_texts.append(row_text)
        if table_texts:
            paragraphs.append("\n".join(table_texts))

    return "\n\n".join(paragraphs)


def _extract_csv(file_path: Path) -> str:
    import csv

    rows: list[str] = []
    with open(file_path, "r", encoding="utf-8", newline="") as f:
        reader = csv.reader(f)
        for row in reader:
            rows.append(" | ".join(row))
    return "\n".join(rows)

def load(
    source: Literal["slack", "github", "notion", "jira", "teams", "confluence", "linear", "zendesk", "google_drive", "documents"],
    path: str | None = None,
    live: bool = False,
    **kwargs: Any,
) -> list[Document]:
    """Load data from a source.

    Parameters
    ----------
    source : {"slack", "github", "notion", "jira", "teams", "confluence", "linear", "zendesk", "google_drive", "documents"}
        The data source type.
    path : str | None
        Path to local export directory. Required for local mode.
    live : bool
        If True, fetch data via live API instead of local export.
    **kwargs
        Additional arguments passed to the loader.

    Returns
    -------
    list[Document]
        LlamaIndex Document objects with metadata["source_type"] set.

    Raises
    ------
    NotImplementedError
        For sources not yet supported (github, notion, jira).
    FileNotFoundError
        If path is provided but does not exist.
    RuntimeError
        If live mode is selected but required env vars are missing.
    """
    if source == "slack":
        if live:
            return _load_slack_live()
        if path is None:
            raise ValueError("path is required for local Slack export loading")
        return _load_slack_export(path, **kwargs)

    if source == "github":
        if live:
            return _load_github()
        raise NotImplementedError(
            "Local GitHub export loading is not supported. Use --live to fetch from the API."
        )

    if source == "notion":
        if live:
            return _load_notion()
        raise NotImplementedError(
            "Local Notion export loading is not supported. Use --live to fetch from the API."
        )

    if source == "jira":
        if live:
            return _load_jira()
        raise NotImplementedError(
            "Local Jira export loading is not supported. Use --live to fetch from the API."
        )

    if source == "teams":
        if live:
            return _load_teams()
        raise NotImplementedError(
            "Local Teams export loading is not supported. Use --live to fetch from the API."
        )

    if source == "confluence":
        if live:
            return _load_confluence()
        raise NotImplementedError(
            "Local Confluence export loading is not supported. Use --live to fetch from the API."
        )

    if source == "linear":
        if live:
            return _load_linear()
        raise NotImplementedError(
            "Local Linear export loading is not supported. Use --live to fetch from the API."
        )

    if source == "zendesk":
        if live:
            return _load_zendesk()
        raise NotImplementedError(
            "Local Zendesk export loading is not supported. Use --live to fetch from the API."
        )

    if source == "google_drive":
        if live:
            return _load_google_drive()
        raise NotImplementedError(
            "Local Google Drive export loading is not supported. Use --live to fetch from the API."
        )

    if source == "documents":
        if path is None:
            raise ValueError("path is required for local document loading")
        return _load_documents(path)

    raise NotImplementedError(f"Source '{source}' is coming soon. Only 'slack', 'github', 'notion', 'jira', 'teams', 'confluence', 'linear', 'zendesk', 'google_drive', and 'documents' are supported.")


if __name__ == "__main__":
    # Test local export
    docs = load("slack", path="data/sample")
    print(f"Loaded {len(docs)} documents from local export\n")

    for i, doc in enumerate(docs[:3]):
        print(f"--- Document {i + 1} ---")
        print(f"ID:   {doc.id_}")
        print(f"Text: {doc.text[:120]}...")
        print("Metadata:")
        for k, v in doc.metadata.items():
            print(f"  {k}: {v}")
        print()
